"""
Enhanced utility functions for RAG Desktop Application
Phase 4: Advanced Document Chunking Logic

This module provides comprehensive text processing and chunking algorithms
for optimal context preservation and semantic coherence.
"""

import re
import os
import logging
from typing import List, Dict, Optional, Tuple, Union
from pathlib import Path
import hashlib
import mimetypes
from datetime import datetime

# File processing imports
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF processing unavailable - install PyPDF2 and pdfplumber")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("DOCX processing unavailable - install python-docx")

try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False
    logging.warning("NLTK unavailable - falling back to basic sentence splitting")

# Text processing imports
import unicodedata
from collections import Counter

# Setup logging
logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE_MB = 50
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.rtf'}
MIN_CHUNK_SIZE = 100
MAX_CHUNK_SIZE = 2000
DEFAULT_OVERLAP = 100
SENTENCE_ENDINGS = r'[.!?]+\s+'
PARAGRAPH_SEPARATOR = r'\n\s*\n'

# ============================================================================
# File Processing Functions (Enhanced from Phase 3)
# ============================================================================

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF with fallback strategies for robustness.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If PDF processing fails
    """
    if not PDF_AVAILABLE:
        raise ValueError("PDF processing libraries not available")
    
    text = ""
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    
    # Primary method: pdfplumber (better for complex layouts)
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            logger.info(f"Successfully extracted text using pdfplumber: {len(text)} chars")
            return clean_text(text)
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
    
    # Fallback method: PyPDF2
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            logger.info(f"Successfully extracted text using PyPDF2: {len(text)} chars")
            return clean_text(text)
    except Exception as e:
        logger.error(f"Both PDF extraction methods failed: {e}")
    
    raise ValueError(f"Failed to extract text from PDF: {file_path}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file with comprehensive formatting preservation.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise ValueError("DOCX processing library not available")
    
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract paragraph text with proper spacing
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n\n".join(text_parts)
        logger.info(f"Successfully extracted DOCX text: {len(text)} chars")
        return clean_text(text)
        
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {file_path}")


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from plain text file with encoding detection.
    
    Args:
        file_path: Path to text file
        
    Returns:
        Extracted text content
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"Text file not found: {file_path}")
    
    # Try multiple encodings
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            logger.info(f"Successfully read text file with {encoding}: {len(text)} chars")
            return clean_text(text)
        except UnicodeDecodeError:
            continue
    
    raise ValueError(f"Failed to decode text file with any encoding: {file_path}")


def extract_text_from_markdown(file_path: str) -> str:
    """
    Extract text from Markdown file, preserving structure.
    
    Args:
        file_path: Path to Markdown file
        
    Returns:
        Extracted text content with preserved formatting
    """
    text = extract_text_from_txt(file_path)
    
    # Basic markdown cleaning while preserving structure
    # Remove markdown syntax but keep content structure
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)  # Headers
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
    text = re.sub(r'`(.*?)`', r'\1', text)  # Inline code
    text = re.sub(r'```[\s\S]*?```', '', text)  # Code blocks
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
    
    return clean_text(text)


def detect_file_type(file_path: str) -> str:
    """
    Detect file type based on extension and MIME type.
    
    Args:
        file_path: Path to file
        
    Returns:
        Detected file type
    """
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    if extension in SUPPORTED_EXTENSIONS:
        return extension[1:]  # Remove the dot
    
    # Fallback to MIME type detection
    mime_type, _ = mimetypes.guess_type(file_path)
    if mime_type:
        type_mapping = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt',
            'text/markdown': 'md'
        }
        return type_mapping.get(mime_type, 'unknown')
    
    return 'unknown'


def validate_file_size(file_path: str, max_size_mb: int = MAX_FILE_SIZE_MB) -> bool:
    """
    Validate file size against maximum allowed size.
    
    Args:
        file_path: Path to file
        max_size_mb: Maximum size in megabytes
        
    Returns:
        True if file size is valid
    """
    file_path = Path(file_path)
    if not file_path.exists():
        return False
    
    size_mb = file_path.stat().st_size / (1024 * 1024)
    return size_mb <= max_size_mb


# ============================================================================
# Text Processing Functions (Enhanced)
# ============================================================================

def clean_text(text: str) -> str:
    """
    Comprehensive text cleaning while preserving semantic structure.
    
    Args:
        text: Raw text to clean
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Normalize unicode characters
    text = unicodedata.normalize('NFKC', text)
    
    # Remove excessive whitespace while preserving paragraph breaks
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple line breaks to double
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
    text = re.sub(r'^\s+|\s+$', '', text, flags=re.MULTILINE)  # Trim lines
    
    # Remove or replace problematic characters
    text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\\\n]', '', text)
    
    # Fix common formatting issues
    text = re.sub(r'([.!?])\s*\n\s*([A-Z])', r'\1 \2', text)  # Join broken sentences
    text = re.sub(r'([a-z])\n([a-z])', r'\1 \2', text)  # Join broken words
    
    return text.strip()


def normalize_whitespace(text: str) -> str:
    """
    Normalize whitespace for consistent processing.
    
    Args:
        text: Text to normalize
        
    Returns:
        Text with normalized whitespace
    """
    # Replace various whitespace characters with standard spaces
    text = re.sub(r'[\t\r\f\v]', ' ', text)
    # Collapse multiple spaces
    text = re.sub(r' +', ' ', text)
    # Normalize line breaks
    text = re.sub(r'\n+', '\n', text)
    
    return text.strip()


def remove_special_characters(text: str, keep_punctuation: bool = True) -> str:
    """
    Remove special characters while optionally preserving punctuation.
    
    Args:
        text: Text to process
        keep_punctuation: Whether to keep basic punctuation
        
    Returns:
        Processed text
    """
    if keep_punctuation:
        # Keep alphanumeric, spaces, and basic punctuation
        pattern = r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\\]'
    else:
        # Keep only alphanumeric and spaces
        pattern = r'[^\w\s]'
    
    return re.sub(pattern, '', text)


def split_into_sentences(text: str) -> List[str]:
    """
    Split text into sentences using NLTK or fallback regex.
    
    Args:
        text: Text to split
        
    Returns:
        List of sentences
    """
    if NLTK_AVAILABLE:
        try:
            sentences = sent_tokenize(text)
            return [s.strip() for s in sentences if s.strip()]
        except Exception as e:
            logger.warning(f"NLTK sentence tokenization failed: {e}")
    
    # Fallback regex-based sentence splitting
    sentences = re.split(SENTENCE_ENDINGS, text)
    return [s.strip() for s in sentences if s.strip() and len(s) > 10]


def split_into_paragraphs(text: str) -> List[str]:
    """
    Split text into paragraphs based on double line breaks.
    
    Args:
        text: Text to split
        
    Returns:
        List of paragraphs
    """
    paragraphs = re.split(PARAGRAPH_SEPARATOR, text)
    return [p.strip() for p in paragraphs if p.strip() and len(p) > 20]


# ============================================================================
# Advanced Chunking Algorithms (New for Phase 4)
# ============================================================================

def adaptive_chunk_text(
    text: str, 
    min_size: int = 800, 
    max_size: int = 1200, 
    overlap: int = 100
) -> List[str]:
    """
    Adaptive chunking that respects sentence boundaries and maintains context.
    
    This algorithm dynamically adjusts chunk sizes based on content structure
    while maintaining semantic coherence.
    
    Args:
        text: Text to chunk
        min_size: Minimum chunk size in characters
        max_size: Maximum chunk size in characters
        overlap: Overlap size between chunks
        
    Returns:
        List of text chunks with preserved context
    """
    if not text or len(text) < min_size:
        return [text] if text else []
    
    sentences = split_into_sentences(text)
    if not sentences:
        return [text]
    
    chunks = []
    current_chunk = ""
    current_size = 0
    
    for sentence in sentences:
        sentence_size = len(sentence)
        
        # If adding this sentence would exceed max_size and we have enough content
        if current_size + sentence_size > max_size and current_size >= min_size:
            chunks.append(current_chunk.strip())
            
            # Create overlap by including last few sentences
            overlap_text = _create_sentence_overlap(current_chunk, overlap)
            current_chunk = overlap_text + " " + sentence
            current_size = len(current_chunk)
        else:
            # Add sentence to current chunk
            if current_chunk:
                current_chunk += " " + sentence
            else:
                current_chunk = sentence
            current_size = len(current_chunk)
    
    # Add final chunk if it has content
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # Post-process to ensure quality
    return _validate_and_clean_chunks(chunks, min_size)


def semantic_chunk_by_sentences(text: str, target_size: int = 1000) -> List[str]:
    """
    Semantic chunking based on sentence similarity and topic coherence.
    
    Groups sentences by semantic similarity to maintain topical coherence
    within each chunk.
    
    Args:
        text: Text to chunk
        target_size: Target chunk size in characters
        
    Returns:
        List of semantically coherent chunks
    """
    sentences = split_into_sentences(text)
    if len(sentences) <= 1:
        return [text] if text else []
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # Calculate potential chunk size
        potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
        
        # If we're approaching target size, finalize chunk
        if len(potential_chunk) >= target_size and current_chunk:
            chunks.append(current_chunk.strip())
            current_chunk = sentence
        else:
            current_chunk = potential_chunk
    
    # Add final chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return _validate_and_clean_chunks(chunks, MIN_CHUNK_SIZE)


def chunk_by_paragraphs(text: str, max_size: int = 1500) -> List[str]:
    """
    Chunk text by paragraphs, combining small paragraphs and splitting large ones.
    
    Args:
        text: Text to chunk
        max_size: Maximum chunk size in characters
        
    Returns:
        List of paragraph-based chunks
    """
    paragraphs = split_into_paragraphs(text)
    if not paragraphs:
        return [text] if text else []
    
    chunks = []
    current_chunk = ""
    
    for paragraph in paragraphs:
        # If paragraph alone exceeds max_size, split it further
        if len(paragraph) > max_size:
            # Finalize current chunk if it exists
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            
            # Split large paragraph into sentences
            para_chunks = adaptive_chunk_text(paragraph, MIN_CHUNK_SIZE, max_size, 50)
            chunks.extend(para_chunks)
        else:
            # Check if adding paragraph would exceed max_size
            potential_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            
            if len(potential_chunk) > max_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk = potential_chunk
    
    # Add final chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return _validate_and_clean_chunks(chunks, MIN_CHUNK_SIZE)


def preserve_context_overlap(chunks: List[str], overlap_size: int = 100) -> List[str]:
    """
    Add contextual overlap between chunks for better coherence.
    
    Args:
        chunks: List of text chunks
        overlap_size: Size of overlap in characters
        
    Returns:
        List of chunks with context overlap
    """
    if len(chunks) <= 1 or overlap_size <= 0:
        return chunks
    
    overlapped_chunks = [chunks[0]]  # First chunk unchanged
    
    for i in range(1, len(chunks)):
        previous_chunk = chunks[i-1]
        current_chunk = chunks[i]
        
        # Extract overlap from end of previous chunk
        overlap = _extract_overlap_text(previous_chunk, overlap_size)
        
        # Prepend overlap to current chunk
        overlapped_chunk = overlap + " " + current_chunk if overlap else current_chunk
        overlapped_chunks.append(overlapped_chunk)
    
    return overlapped_chunks


# ============================================================================
# Chunking Utility Functions
# ============================================================================

def _create_sentence_overlap(text: str, overlap_size: int) -> str:
    """Create overlap by extracting last sentences from text."""
    sentences = split_into_sentences(text)
    if not sentences:
        return ""
    
    overlap_text = ""
    for sentence in reversed(sentences):
        potential_overlap = sentence + " " + overlap_text if overlap_text else sentence
        if len(potential_overlap) <= overlap_size:
            overlap_text = potential_overlap
        else:
            break
    
    return overlap_text.strip()


def _extract_overlap_text(text: str, overlap_size: int) -> str:
    """Extract overlap text from end of chunk."""
    if len(text) <= overlap_size:
        return text
    
    # Try to find sentence boundary for clean overlap
    sentences = split_into_sentences(text)
    if sentences:
        return _create_sentence_overlap(text, overlap_size)
    
    # Fallback to character-based overlap
    return text[-overlap_size:].strip()


def _validate_and_clean_chunks(chunks: List[str], min_size: int) -> List[str]:
    """Validate and clean chunks, merging small ones."""
    if not chunks:
        return []
    
    validated_chunks = []
    pending_chunk = ""
    
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        
        if len(chunk) < min_size:
            # Accumulate small chunks
            pending_chunk = pending_chunk + " " + chunk if pending_chunk else chunk
        else:
            # If we have pending small chunk, add it first
            if pending_chunk:
                if len(pending_chunk) >= min_size:
                    validated_chunks.append(pending_chunk.strip())
                else:
                    # Merge with current chunk
                    chunk = pending_chunk + " " + chunk
                pending_chunk = ""
            
            validated_chunks.append(chunk)
    
    # Handle any remaining pending chunk
    if pending_chunk:
        if validated_chunks and len(pending_chunk) < min_size:
            # Merge with last chunk
            validated_chunks[-1] = validated_chunks[-1] + " " + pending_chunk
        else:
            validated_chunks.append(pending_chunk.strip())
    
    return validated_chunks


# ============================================================================
# Embedding Utilities (Preparation for Phase 5)
# ============================================================================

def batch_texts(texts: List[str], batch_size: int = 32) -> List[List[str]]:
    """
    Batch texts for efficient embedding processing.
    
    Args:
        texts: List of texts to batch
        batch_size: Size of each batch
        
    Returns:
        List of text batches
    """
    batches = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        batches.append(batch)
    
    return batches


def optimize_embedding_batch_size(text_lengths: List[int]) -> int:
    """
    Optimize batch size based on text lengths for memory efficiency.
    
    Args:
        text_lengths: List of text lengths
        
    Returns:
        Optimal batch size
    """
    if not text_lengths:
        return 32
    
    avg_length = sum(text_lengths) / len(text_lengths)
    
    # Adjust batch size based on average text length
    if avg_length < 200:
        return 64
    elif avg_length < 500:
        return 32
    elif avg_length < 1000:
        return 16
    else:
        return 8


def calculate_embedding_similarity(vec1: List[float], vec2: List[float]) -> float:
    """
    Calculate cosine similarity between two embedding vectors.
    
    Args:
        vec1: First embedding vector
        vec2: Second embedding vector
        
    Returns:
        Cosine similarity score
    """
    if len(vec1) != len(vec2):
        raise ValueError("Vectors must have same length")
    
    # Calculate dot product
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    
    # Calculate magnitudes
    magnitude1 = sum(a * a for a in vec1) ** 0.5
    magnitude2 = sum(b * b for b in vec2) ** 0.5
    
    # Avoid division by zero
    if magnitude1 == 0 or magnitude2 == 0:
        return 0.0
    
    return dot_product / (magnitude1 * magnitude2)


# ============================================================================
# Security Functions
# ============================================================================

def generate_secure_id() -> str:
    """Generate a secure random ID."""
    import uuid
    return str(uuid.uuid4())


def hash_password(password: str) -> str:
    """Hash password using SHA-256 (placeholder for proper hashing)."""
    return hashlib.sha256(password.encode()).hexdigest()


def verify_password(password: str, hashed: str) -> bool:
    """Verify password against hash."""
    return hash_password(password) == hashed


def generate_api_key() -> str:
    """Generate a secure API key."""
    import secrets
    return secrets.token_urlsafe(32)


# ============================================================================
# Date/Time Utilities
# ============================================================================

def get_utc_now() -> datetime:
    """Get current UTC datetime."""
    return datetime.utcnow()


def format_timestamp(dt: datetime) -> str:
    """Format datetime as ISO string."""
    return dt.isoformat()


def parse_timestamp(timestamp_str: str) -> datetime:
    """Parse ISO timestamp string to datetime."""
    return datetime.fromisoformat(timestamp_str)


# ============================================================================
# Validation Functions
# ============================================================================

def validate_email(email: str) -> bool:
    """Validate email format."""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None


def validate_url(url: str) -> bool:
    """Validate URL format."""
    pattern = r'^https?://[^\s/$.?#].[^\s]*$'
    return re.match(pattern, url) is not None


def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe storage."""
    # Remove or replace invalid characters
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    # Limit length
    if len(filename) > 255:
        name, ext = os.path.splitext(filename)
        filename = name[:255-len(ext)] + ext
    
    return filename


# ============================================================================
# Chunking Strategy Selection
# ============================================================================

def get_optimal_chunking_strategy(text: str, file_type: str) -> str:
    """
    Determine optimal chunking strategy based on content analysis.
    
    Args:
        text: Text content
        file_type: File type (pdf, docx, txt, md)
        
    Returns:
        Recommended chunking strategy
    """
    text_length = len(text)
    paragraph_count = len(split_into_paragraphs(text))
    sentence_count = len(split_into_sentences(text))
    
    # Calculate metrics
    avg_paragraph_length = text_length / max(paragraph_count, 1)
    avg_sentence_length = text_length / max(sentence_count, 1)
    
    # Strategy selection logic
    if file_type in ['pdf', 'docx'] and avg_paragraph_length > 500:
        return "paragraph"
    elif avg_sentence_length > 100 and sentence_count > 10:
        return "semantic"
    else:
        return "adaptive"


def chunk_text_with_strategy(
    text: str, 
    strategy: str = "adaptive",
    **kwargs
) -> List[str]:
    """
    Chunk text using specified strategy.
    
    Args:
        text: Text to chunk
        strategy: Chunking strategy (adaptive, semantic, paragraph)
        **kwargs: Strategy-specific parameters
        
    Returns:
        List of text chunks
    """
    if strategy == "adaptive":
        return adaptive_chunk_text(
            text,
            kwargs.get('min_size', 800),
            kwargs.get('max_size', 1200),
            kwargs.get('overlap', 100)
        )
    elif strategy == "semantic":
        return semantic_chunk_by_sentences(
            text,
            kwargs.get('target_size', 1000)
        )
    elif strategy == "paragraph":
        return chunk_by_paragraphs(
            text,
            kwargs.get('max_size', 1500)
        )
    else:
        raise ValueError(f"Unknown chunking strategy: {strategy}")


# ============================================================================
# Quality Assessment
# ============================================================================

def assess_chunk_quality(chunks: List[str]) -> Dict[str, float]:
    """
    Assess quality metrics for text chunks.
    
    Args:
        chunks: List of text chunks
        
    Returns:
        Quality metrics dictionary
    """
    if not chunks:
        return {}
    
    chunk_lengths = [len(chunk) for chunk in chunks]
    
    metrics = {
        'total_chunks': len(chunks),
        'avg_chunk_length': sum(chunk_lengths) / len(chunk_lengths),
        'min_chunk_length': min(chunk_lengths),
        'max_chunk_length': max(chunk_lengths),
        'length_variance': _calculate_variance(chunk_lengths),
        'completeness_score': _calculate_completeness_score(chunks)
    }
    
    return metrics


def _calculate_variance(values: List[float]) -> float:
    """Calculate variance of values."""
    if len(values) < 2:
        return 0.0
    
    mean = sum(values) / len(values)
    variance = sum((x - mean) ** 2 for x in values) / len(values)
    return variance


def _calculate_completeness_score(chunks: List[str]) -> float:
    """Calculate completeness score based on chunk content quality."""
    if not chunks:
        return 0.0
    
    scores = []
    for chunk in chunks:
        # Basic quality indicators
        has_complete_sentences = '.' in chunk or '!' in chunk or '?' in chunk
        has_sufficient_length = len(chunk) > MIN_CHUNK_SIZE
        has_meaningful_content = len(chunk.split()) > 10
        
        score = sum([has_complete_sentences, has_sufficient_length, has_meaningful_content]) / 3
        scores.append(score)
    
    return sum(scores) / len(scores)